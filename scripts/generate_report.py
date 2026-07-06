"""
Generate comprehensive .docx report for Furniture Sales Revenue & Profit Optimization.
Matches the format of MauBaoCao.docx template while including ALL analytics results.
"""
import sys
import os

sys.stdout.reconfigure(encoding="utf-8")

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import csv

# ── Paths ──────────────────────────────────────────────────────────────
PROJECT = Path(__file__).resolve().parent.parent
REPORTS = PROJECT / "reports"
TABLES = REPORTS / "tables"
FIGURES = REPORTS / "figures"
OUTPUT = PROJECT / "BaoCao_PhanTichDuLieu.docx"

# ── Helpers ────────────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Set background shading on a cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a table with Light Grid Accent 1 style and formatted content."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.size = Pt(10)
                r.font.bold = True

    # Data rows
    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(val)
            for p in cell.paragraphs:
                # Right-align numeric columns (heuristic: starts with $, -, digit, or contains %)
                txt = str(val).strip()
                if txt and (txt[0] in "$-0123456789" or "%" in txt or txt == "N/A"):
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    r.font.size = Pt(9)

    # Apply widths if specified
    if col_widths:
        for row in table.rows:
            for j, w in enumerate(col_widths):
                if j < len(row.cells):
                    row.cells[j].width = Cm(w)

    return table


def read_csv_data(filename):
    """Read CSV file and return headers, rows."""
    path = TABLES / filename
    if not path.exists():
        return [], []
    with open(path, "r", encoding="utf-8-sig") as f:
        reader = csv.reader(f)
        headers = next(reader, [])
        rows = list(reader)
    return headers, rows


def add_figure(doc, filename, caption, width=Inches(5.5)):
    """Add a figure with caption."""
    path = FIGURES / filename
    if path.exists():
        doc.add_picture(str(path), width=width)
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.style = "Normal"
        for r in cap.runs:
            r.font.size = Pt(9)
            r.font.italic = True
    else:
        doc.add_paragraph(f"[Hình không tìm thấy: {filename}]")


def add_observation(doc, text):
    """Add an observation/fact paragraph with bold prefix."""
    p = doc.add_paragraph()
    run_label = p.add_run("Sự thật quan sát: ")
    run_label.bold = True
    run_label.font.size = Pt(11)
    run_label.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    run_text = p.add_run(text)
    run_text.font.size = Pt(11)


def add_hypothesis(doc, text):
    """Add an evidence-backed hypothesis paragraph."""
    p = doc.add_paragraph()
    run_label = p.add_run("Giả thuyết có bằng chứng: ")
    run_label.bold = True
    run_label.font.size = Pt(11)
    run_label.font.color.rgb = RGBColor(0x8B, 0x6D, 0x0B)
    run_text = p.add_run(text)
    run_text.font.size = Pt(11)


def add_implication(doc, text):
    """Add a business implication paragraph."""
    p = doc.add_paragraph()
    run_label = p.add_run("Hàm ý kinh doanh: ")
    run_label.bold = True
    run_label.font.size = Pt(11)
    run_label.font.color.rgb = RGBColor(0x0B, 0x6E, 0x0B)
    run_text = p.add_run(text)
    run_text.font.size = Pt(11)


def add_spacer(doc):
    """Add an empty paragraph as spacer."""
    doc.add_paragraph("")


# ══════════════════════════════════════════════════════════════════════
#  MAIN REPORT GENERATION
# ══════════════════════════════════════════════════════════════════════

def generate_report():
    doc = Document()

    # ── Page setup ─────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)

    # ══════════════════════════════════════════════════════════════
    #  COVER PAGE
    # ══════════════════════════════════════════════════════════════
    for _ in range(3):
        doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("TRƯỜNG ĐẠI HỌC NGOẠI NGỮ - TIN HỌC TP.HCM")
    r.font.size = Pt(14)
    r.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Khoa Công Nghệ Thông Tin")
    r.font.size = Pt(12)
    r.bold = True

    for _ in range(4):
        doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Đề tài")
    r.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "Phân tích dữ liệu bán hàng đồ gỗ\ntối ưu hoá doanh thu và lợi nhuận"
    )
    r.font.size = Pt(16)
    r.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    for _ in range(5):
        doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("GVHD: [Tên Giáo Viên Hướng Dẫn]")
    r.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Tháng 7 Năm 2026")
    r.font.size = Pt(11)
    r.bold = True

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    #  TABLE OF CONTENTS (placeholder)
    # ══════════════════════════════════════════════════════════════
    h = doc.add_heading("Mục lục", level=1)
    p = doc.add_paragraph(
        "I. Giới thiệu\n"
        "II. Nội dung dữ liệu\n"
        "III. Phương pháp phân tích\n"
        "IV. Kết quả phân tích\n"
        "    4.1. Tổng quan KPI chính\n"
        "    4.2. Hiệu quả điều hành theo thời gian (Nhóm A)\n"
        "    4.3. Hiệu quả sản phẩm và danh mục (Nhóm B)\n"
        "    4.4. Rủi ro discount và pricing (Nhóm C)\n"
        "    4.5. Hiệu quả địa lý (Nhóm D)\n"
        "    4.6. Khách hàng và segment (Nhóm E)\n"
        "    4.7. Fulfillment context (Nhóm F)\n"
        "V. Khuyến nghị hành động\n"
        "VI. Tổng quan và kết luận\n"
        "VII. Phụ lục"
    )
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    #  I. GIỚI THIỆU
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("I. Giới thiệu", level=1)

    doc.add_paragraph(
        "Bộ dữ liệu phân tích bán hàng đồ gỗ (Furniture Sales Dataset) chứa thông tin "
        "chi tiết về 2.121 dòng giao dịch thuộc 1.764 đơn hàng của 707 khách hàng, bao gồm "
        "thông tin sản phẩm, thời gian, địa điểm, khách hàng, và doanh thu / lợi nhuận. "
        "Dữ liệu bao phủ giai đoạn từ tháng 1 năm 2014 đến tháng 12 năm 2017, chỉ thuộc "
        "danh mục Furniture (đồ nội thất)."
    )

    doc.add_paragraph(
        "Mục tiêu của đề tài là phân tích khám phá (EDA) bộ dữ liệu này nhằm hiểu rõ:"
    )

    bullets = [
        "Doanh thu và lợi nhuận được tạo ra từ những sản phẩm, vùng, và nhóm khách hàng nào",
        "Mối quan hệ giữa mức giảm giá (discount) và lợi nhuận",
        "Các yếu tố rủi ro chính ảnh hưởng đến chất lượng lợi nhuận",
        "Xác định sản phẩm, vùng, và nhóm khách hàng cần ưu tiên sửa chữa hoặc mở rộng",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_paragraph(
        "Để thực hiện đề tài, nhóm sử dụng các công cụ xử lý dữ liệu (Python, Pandas, NumPy) "
        "để làm sạch, biến đổi và tính toán các chỉ tiêu kinh doanh. Kết quả được trình bày "
        "qua các biểu đồ chuyên nghiệp (matplotlib, plotly) và bảng phân tích chi tiết. "
        "Phân tích tuân thủ nguyên tắc: phân biệt sự thật quan sát được, giả thuyết có "
        "bằng chứng, và khuyến nghị hành động — không trình bày tương quan như nhân quả."
    )

    doc.add_paragraph(
        "Câu hỏi quyết định chính: Nhà bán lẻ nên thay đổi quyết định sản phẩm, địa lý, "
        "và giảm giá như thế nào để cải thiện lợi nhuận Furniture mà vẫn duy trì hoặc tăng "
        "doanh thu bền vững?"
    )

    add_spacer(doc)

    # ══════════════════════════════════════════════════════════════
    #  II. NỘI DUNG DỮ LIỆU
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("II. Nội dung dữ liệu", level=1)

    doc.add_paragraph("Tên bộ dữ liệu: Furniture Sales Revenue & Profit Optimization Dataset")
    doc.add_paragraph(
        "Nguồn dữ liệu: Dữ liệu bán hàng lịch sử từ hệ thống bán lẻ, "
        "phân loại và xác nhận là danh mục Furniture."
    )
    doc.add_paragraph(
        "Cấu trúc: 2.121 dòng dữ liệu, 21 cột thuộc tính. Các nhóm biến chính:"
    )

    # Variable groups table
    add_styled_table(
        doc,
        ["Nhóm biến", "Ý nghĩa"],
        [
            ["Định danh", "Row ID, Order ID, Customer ID, Product ID"],
            ["Thời gian", "Order Date (2014-01-06 đến 2017-12-30), Ship Date"],
            ["Địa lý", "4 Regions, 48 States, 371 Cities, Postal Code"],
            ["Sản phẩm", "Category (Furniture), Sub-Category: Chairs, Tables, Bookcases, Furnishings"],
            ["Khách hàng", "Customer Name, Segment: Consumer, Corporate, Home Office"],
            ["Tài chính", "Sales (doanh thu), Profit (lợi nhuận), Discount (giảm giá), Quantity (số lượng)"],
            ["Vận chuyển", "Ship Mode: Standard Class, Second Class, First Class, Same Day"],
        ],
    )

    add_spacer(doc)
    doc.add_paragraph("Các chỉ tiêu chính tính từ dữ liệu gốc:")

    metric_bullets = [
        "Weighted profit margin = Tổng lợi nhuận / Tổng doanh thu (biên có trọng số, không trung bình đơn giản)",
        "Loss-line rate = Số dòng lỗ / Tổng dòng giao dịch",
        "Average Order Value (AOV) = Tổng doanh thu / Số đơn hàng duy nhất (nunique Order ID)",
        "Profit per order = Tổng lợi nhuận / Số đơn hàng duy nhất",
        "Sales growth = (Doanh thu kỳ t − Doanh thu kỳ t−1) / Doanh thu kỳ t−1",
        "Discount bands: 0%, 1-10%, 11-20%, 21-30%, >30% (không chồng lấn)",
    ]
    for b in metric_bullets:
        doc.add_paragraph(b, style="List Bullet")

    add_spacer(doc)

    # ══════════════════════════════════════════════════════════════
    #  III. PHƯƠNG PHÁP PHÂN TÍCH
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("III. Phương pháp phân tích", level=1)

    doc.add_paragraph("Quy trình phân tích được thực hiện theo các bước sau:")

    steps = [
        (
            "Bước 1 — Kiểm tra chất lượng dữ liệu: ",
            "Kiểm tra kích thước (2.121 dòng, 21 cột), kiểu dữ liệu, giá trị thiếu, "
            "bản ghi trùng lặp, Row ID duy nhất, danh mục Furniture, và ràng buộc kinh doanh "
            "(Sales ≥ 0, Quantity > 0, 0 ≤ Discount ≤ 1, Ship Date ≥ Order Date).",
        ),
        (
            "Bước 2 — Xác nhận baseline: ",
            "Tính tổng doanh thu ($741.999,80), lợi nhuận ($18.451,27), số đơn hàng (1.764), "
            "khách hàng (707) để đối sánh với giá trị kỳ vọng trong CONTEXT.md.",
        ),
        (
            "Bước 3 — Biến đổi và feature engineering: ",
            "Tạo các cột phái sinh: order_year, order_quarter, order_month, shipping_days, "
            "profit_margin, is_loss_line, discount_band, is_repeat_customer, customer_order_count, v.v.",
        ),
        (
            "Bước 4 — Phân tích thời gian (Nhóm A): ",
            "Phân tích doanh thu, lợi nhuận, biên theo tháng, quý, năm để nhận diện xu hướng "
            "và sự phân kỳ giữa tăng trưởng doanh thu và lợi nhuận.",
        ),
        (
            "Bước 5 — Phân tích sản phẩm (Nhóm B): ",
            "So sánh 4 sub-category, phân loại sản phẩm theo ma trận Grow/Repair/Scale/Rationalize, "
            "xác định top sản phẩm lỗ và sản phẩm lợi nhuận chưa đại diện.",
        ),
        (
            "Bước 6 — Phân tích giảm giá (Nhóm C): ",
            "Phân tích mối quan hệ giữa mức discount và các chỉ tiêu theo band và exact level. "
            "Xác định ngưỡng discount đầu tiên có profit âm và các pattern lỗ không do discount.",
        ),
        (
            "Bước 7 — Phân tích địa lý (Nhóm D): ",
            "Phân tích doanh thu, lợi nhuận, discount theo 4 vùng và 48 tiểu bang. "
            "Tính priority score cho mỗi tiểu bang. Heatmap region × sub-category.",
        ),
        (
            "Bước 8 — Phân tích khách hàng (Nhóm E): ",
            "Phân tích segment performance, repeat customer behavior, và xác định "
            "khách hàng có doanh thu cao nhưng lợi nhuận thấp/âm.",
        ),
        (
            "Bước 9 — Fulfillment context (Nhóm F): ",
            "Phân tích ship mode và shipping days theo doanh thu, lợi nhuận, loss rate. "
            "Mô tả quan sát, không quy kết nguyên nhân do thiếu freight-cost.",
        ),
        (
            "Bước 10 — Tổng hợp và khuyến nghị: ",
            "Xác định các scope cần ưu tiên, cùng bằng chứng cụ thể, hành động đề xuất, "
            "KPI theo dõi, và hạn chế. Mỗi khuyến nghị gắn evidence, không kết luận nhân quả.",
        ),
    ]
    for label, desc in steps:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(label)
        r.bold = True
        r.font.size = Pt(11)
        r2 = p.add_run(desc)
        r2.font.size = Pt(11)

    add_spacer(doc)

    # Data quality result
    doc.add_heading("Kết quả kiểm tra chất lượng dữ liệu", level=2)

    add_styled_table(
        doc,
        ["Chỉ số", "Giá trị"],
        [
            ["Số dòng", "2.121"],
            ["Số cột", "21"],
            ["Giá trị thiếu", "0"],
            ["Dòng trùng lặp hoàn toàn", "0"],
            ["Row ID trùng lặp", "0"],
            ["Đơn hàng duy nhất (nunique Order ID)", "1.764"],
            ["Khách hàng duy nhất (nunique Customer ID)", "707"],
            ["Sản phẩm duy nhất (Product ID)", "375"],
            ["Tiểu bang", "48"],
            ["Phủ sóng Order Date", "2014-01-06 đến 2017-12-30"],
            ["Phủ sóng Ship Date", "2014-01-10 đến 2018-01-05"],
            ["Category coverage", "Furniture only"],
            ["Ship Date < Order Date", "0 (đạt)"],
            ["Tổng doanh thu", "$741.999,80"],
            ["Tổng lợi nhuận", "$18.451,27"],
            ["Dòng lỗ (Profit < 0)", "714"],
        ],
    )

    add_spacer(doc)
    add_observation(
        doc,
        "Dữ liệu đầu vào đạt tất cả ràng buộc bắt buộc: Furniture-only, không có giá trị thiếu, "
        "không có Row ID trùng lặp, ngày giao hàng không sớm hơn ngày đặt hàng, "
        "và các tổng doanh thu/lợi nhuận khớp baseline trong CONTEXT.md. Bảng fact sau biến đổi "
        "giữ nguyên 2.121 dòng với tổng doanh thu $741.999,80 và tổng lợi nhuận $18.451,27."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    #  IV. KẾT QUẢ PHÂN TÍCH
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("IV. Kết quả phân tích", level=1)

    # ── 4.1 KPI Overview ──────────────────────────────────────────
    doc.add_heading("4.1. Tổng quan KPI chính", level=2)
    doc.add_paragraph("Bảng dưới đây tóm tắt các chỉ tiêu chính cho toàn bộ bộ dữ liệu:")

    add_styled_table(
        doc,
        ["KPI", "Giá trị"],
        [
            ["Tổng doanh thu (Total Sales)", "$741.999,80"],
            ["Tổng lợi nhuận (Total Profit)", "$18.451,27"],
            ["Biên lợi nhuận có trọng số", "2,49%"],
            ["Tổng số lượng (Quantity)", "8.028"],
            ["Số đơn hàng duy nhất", "1.764"],
            ["Số khách hàng duy nhất", "707"],
            ["Giá trị đơn hàng trung bình (AOV)", "$420,63"],
            ["Lợi nhuận trung bình/đơn hàng", "$10,46"],
            ["Số dòng lỗ (Profit < 0)", "714 / 2.121"],
            ["Tỷ lệ dòng lỗ", "33,66%"],
            ["Khách hàng có ≥2 đơn (repeat)", "516 / 707 (72,99%)"],
            ["Tỷ lệ đơn từ repeat customers", "89,17%"],
        ],
    )

    add_spacer(doc)
    add_observation(
        doc,
        "Biên lợi nhuận chỉ 2,49% là rất thấp cho ngành bán lẻ nội thất. 33,66% tổng dòng "
        "giao dịch bị lỗ, cho thấy lỗ là vấn đề cấu trúc cần xử lý, không phải ngoại lệ nhỏ."
    )

    doc.add_page_break()

    # ── 4.2 Executive Performance (Group A) ──────────────────────
    doc.add_heading("4.2. Hiệu quả điều hành theo thời gian (Nhóm A)", level=2)

    # A1: Yearly
    doc.add_paragraph().add_run(
        "A1. Doanh thu, lợi nhuận và biên biến động thế nào theo năm?"
    ).bold = True

    add_styled_table(
        doc,
        ["Năm", "Doanh thu", "Lợi nhuận", "Biên %", "Tăng DT", "Tăng LN"],
        [
            ["2014", "$157.192,85", "$5.457,73", "3,47%", "—", "—"],
            ["2015", "$170.518,24", "$3.015,20", "1,77%", "8,48%", "-44,75%"],
            ["2016", "$198.901,44", "$6.959,95", "3,50%", "16,65%", "130,83%"],
            ["2017", "$215.387,27", "$3.018,39", "1,40%", "8,29%", "-56,63%"],
        ],
    )

    add_spacer(doc)
    add_observation(
        doc,
        "Doanh thu tăng liên tục 37% từ 2014-2017, nhưng lợi nhuận không tăng theo tỷ lệ. "
        "Năm 2015 lợi nhuận giảm 44,75%. Năm 2017 đạt doanh thu cao nhất ($215.387) "
        "nhưng biên lợi nhuận thấp nhất (1,40%)."
    )

    # Sales & margin by year chart
    add_figure(doc, "sales_margin_by_year.png", "Hình 1: Doanh thu và biên lợi nhuận theo năm")

    add_spacer(doc)

    # Monthly trend
    doc.add_paragraph().add_run(
        "A1 (tiếp). Xu hướng doanh thu và lợi nhuận theo tháng:"
    ).bold = True

    add_figure(doc, "monthly_sales_profit.png", "Hình 2: Xu hướng doanh thu và lợi nhuận theo tháng (2014-2017)")

    add_spacer(doc)

    # A2-A3: Divergence
    doc.add_paragraph().add_run(
        "A2-A3. Tăng trưởng doanh thu có chuyển thành tăng trưởng lợi nhuận không?"
    ).bold = True

    doc.add_paragraph(
        "Các kỳ dưới đây có doanh thu xếp hạng cao hơn nhiều so với lợi nhuận "
        "(divergence score = sales_rank − profit_rank):"
    )

    add_styled_table(
        doc,
        ["Tháng", "Doanh thu", "Lợi nhuận", "Biên %", "Sales Rank", "Profit Rank", "Divergence"],
        [
            ["2017-10", "$21.884,07", "$-2.526,92", "-11,55%", "12", "47", "35"],
            ["2014-03", "$14.573,96", "$-1.128,65", "-7,74%", "18", "46", "28"],
            ["2014-11", "$21.564,87", "$-297,90", "-1,38%", "13", "38", "25"],
            ["2017-11", "$37.056,72", "$406,06", "1,10%", "1", "21", "20"],
            ["2015-07", "$13.674,42", "$-325,09", "-2,38%", "19", "39", "20"],
            ["2016-03", "$12.801,09", "$-555,27", "-4,34%", "24", "44", "20"],
            ["2017-05", "$16.957,56", "$-72,88", "-0,43%", "15", "34", "19"],
            ["2016-08", "$12.483,23", "$-494,15", "-3,96%", "26", "43", "17"],
        ],
    )

    add_spacer(doc)
    add_implication(
        doc,
        "Cần review khuyến mại, nhóm sản phẩm và bang trong các kỳ doanh thu cao nhưng profit yếu. "
        "Tháng 11/2017 đạt doanh thu #1 ($37.057) nhưng lợi nhuận chỉ #21 ($406), "
        "cho thấy chất lượng doanh thu kém."
    )

    # A4: Loss-line rate
    doc.add_paragraph().add_run(
        "A4. Tỷ lệ dòng giao dịch bị lỗ:"
    ).bold = True

    add_observation(
        doc,
        "Có 714 dòng lỗ (Profit < 0), tương đương 33,66% tổng số dòng giao dịch. "
        "Lỗ không phải ngoại lệ nhỏ — cần worklist kiểm soát giá/discount ở cấp dòng."
    )

    doc.add_page_break()

    # ── 4.3 Product Performance (Group B) ────────────────────────
    doc.add_heading("4.3. Hiệu quả sản phẩm và danh mục (Nhóm B)", level=2)

    # B1: Sub-category
    doc.add_paragraph().add_run(
        "B1. Sub-category nào tạo doanh thu, lợi nhuận và margin?"
    ).bold = True

    add_styled_table(
        doc,
        ["Sub-Category", "Doanh thu", "Lợi nhuận", "Biên %", "Dòng lỗ %"],
        [
            ["Chairs", "$328.449,10", "$26.590,17", "8,10%", "38,09%"],
            ["Tables", "$206.965,53", "$-17.725,48", "-8,56%", "63,64%"],
            ["Bookcases", "$114.880,00", "$-3.472,56", "-3,02%", "47,81%"],
            ["Furnishings", "$91.705,16", "$13.059,14", "14,24%", "17,45%"],
        ],
    )

    add_spacer(doc)
    add_figure(doc, "sub_category_sales_profit.png", "Hình 3: So sánh doanh thu và lợi nhuận theo sub-category")

    add_spacer(doc)
    add_observation(
        doc,
        "Chairs tạo lợi nhuận lớn nhất (+$26.590, biên 8,10%); Furnishings có biên tốt nhất "
        "(14,24%); Tables lỗ lớn nhất (-$17.725, biên -8,56%); Bookcases cũng lỗ (-$3.473, biên -3,02%). "
        "Tables và Bookcases cần sửa chữa lợi nhuận trước khi mở rộng doanh thu."
    )

    add_spacer(doc)

    # B2: High-sales negative-profit products
    doc.add_paragraph().add_run(
        "B2. Sản phẩm nào doanh thu cao nhưng lợi nhuận thấp hoặc âm (nhóm Repair)?"
    ).bold = True

    add_styled_table(
        doc,
        ["Product ID", "Tên sản phẩm", "Sub-Cat", "Doanh thu", "Lợi nhuận", "Biên %"],
        [
            ["FUR-CH-10002024", "HON 5400 Series Task Chairs for Big and Tall", "Chairs", "$21.870,58", "$0,00", "0,00%"],
            ["FUR-BO-10004834", "Riverside Palais Royal Lawyers Bookcase", "Bookcases", "$15.610,97", "$-669,54", "-4,29%"],
            ["FUR-TA-10003473", "Bretford Rectangular Conference Table Tops", "Tables", "$12.995,29", "$-327,23", "-2,52%"],
            ["FUR-BO-10002213", "DMI Eclipse Executive Suite Bookcases", "Bookcases", "$11.046,61", "$90,18", "0,82%"],
            ["FUR-TA-10000198", "Chromcraft Bull-Nose Wood Oval Conference Tables", "Tables", "$9.917,64", "$-2.876,12", "-29,00%"],
            ["FUR-TA-10001889", "Bush Advantage Collection Racetrack Conference Table", "Tables", "$9.544,73", "$-1.934,40", "-20,27%"],
            ["FUR-TA-10001095", "Chromcraft Round Conference Tables", "Tables", "$8.209,06", "$-189,98", "-2,31%"],
            ["FUR-TA-10003954", "Hon 94000 Series Round Tables", "Tables", "$7.404,50", "$-681,21", "-9,20%"],
            ["FUR-TA-10000577", "Bretford CR4500 Series Slim Rectangular Table", "Tables", "$7.242,77", "$-532,76", "-7,36%"],
            ["FUR-TA-10002958", "Bevis Oval Conference Table, Walnut", "Tables", "$6.942,07", "$-856,01", "-12,33%"],
        ],
    )

    add_spacer(doc)
    add_implication(
        doc,
        "Các sản phẩm này thuộc nhóm Repair: doanh thu cao nhưng lợi nhuận thấp/âm. "
        "Cần xem lại giá, mức giảm giá và vùng bán trước khi tiếp tục scale."
    )

    add_spacer(doc)

    # B3: Underrepresented profitable products
    doc.add_paragraph().add_run(
        "B3. Sản phẩm nào có lợi nhuận tốt nhưng chưa đại diện lớn trong doanh thu (nhóm Scale Selectively)?"
    ).bold = True

    add_styled_table(
        doc,
        ["Product ID", "Tên sản phẩm", "Sub-Cat", "Doanh thu", "Lợi nhuận", "Biên %"],
        [
            ["FUR-FU-10002937", "GE 48\" Fluorescent Tube, Cool White Energy Saver", "Furnishings", "$2.699,06", "$1.260,22", "46,69%"],
            ["FUR-FU-10002671", "Electrix 20W Halogen Replacement Bulb", "Furnishings", "$168,84", "$78,26", "46,35%"],
            ["FUR-FU-10001861", "Floodlight Indoor Halogen Bulbs, 60 Watts", "Furnishings", "$434,56", "$197,10", "45,36%"],
            ["FUR-FU-10002045", "Executive Impressions 14\"", "Furnishings", "$377,91", "$166,28", "44,00%"],
            ["FUR-FU-10000409", "GE 4 Foot Flourescent Tube, 40 Watt", "Furnishings", "$98,87", "$42,24", "42,73%"],
        ],
    )

    add_spacer(doc)
    add_implication(
        doc,
        "Đây là nhóm Scale Selectively: biên rất cao (40-47%) nhưng doanh thu nhỏ. "
        "Có thể cải thiện discoverability và cross-sell có mục tiêu thay vì giảm giá đại trà."
    )

    add_spacer(doc)

    # B4: Top loss products
    doc.add_paragraph().add_run(
        "B4. Sản phẩm nào gây lỗ lớn nhất?"
    ).bold = True

    add_styled_table(
        doc,
        ["Tên sản phẩm", "Sub-Cat", "Doanh thu", "Lợi nhuận", "Biên %"],
        [
            ["Chromcraft Bull-Nose Wood Oval Conference Tables", "Tables", "$9.917,64", "$-2.876,12", "-29,00%"],
            ["Bush Advantage Collection Racetrack Conference Table", "Tables", "$9.544,73", "$-1.934,40", "-20,27%"],
            ["Balt Solid Wood Round Tables", "Tables", "$6.518,75", "$-1.201,06", "-18,42%"],
            ["BoxOffice By Design Meeting Room Tables", "Tables", "$1.706,25", "$-1.148,44", "-67,31%"],
            ["Riverside Furniture Oval Coffee Table", "Tables", "$4.446,18", "$-1.147,40", "-25,81%"],
            ["Hon 2090 Pillow Soft Series Chairs", "Chairs", "$5.282,42", "$-989,05", "-18,72%"],
            ["O'Sullivan 4-Shelf Bookcase", "Bookcases", "$2.740,20", "$-975,10", "-35,58%"],
            ["Bretford Just In Time Work Tables", "Tables", "$5.634,90", "$-964,19", "-17,11%"],
            ["Bevis Oval Conference Table, Walnut", "Tables", "$6.942,07", "$-856,01", "-12,33%"],
            ["BPI Conference Tables", "Tables", "$2.241,87", "$-795,97", "-35,50%"],
        ],
    )

    add_spacer(doc)
    add_figure(doc, "top_loss_products.png", "Hình 4: Top 10 sản phẩm gây lỗ lớn nhất")

    add_spacer(doc)
    add_observation(
        doc,
        "9/10 sản phẩm lỗ lớn nhất là Tables, chủ yếu bàn hội nghị. "
        "Tổng lỗ từ 10 sản phẩm này chiếm phần lớn tổng lỗ của sub-category Tables."
    )

    add_spacer(doc)

    # B5: Region x Sub-category
    doc.add_paragraph().add_run(
        "B5. Tổ hợp region × sub-category nào tạo hoặc hủy lợi nhuận?"
    ).bold = True

    add_styled_table(
        doc,
        ["Region", "Sub-Category", "Doanh thu", "Lợi nhuận", "Biên %", "Disc %", "Loss %"],
        [
            ["East", "Tables", "$39.139,81", "$-11.025,38", "-28,17%", "37,37%", "97,50%"],
            ["South", "Tables", "$43.916,19", "$-4.623,06", "-10,53%", "22,25%", "52,94%"],
            ["Central", "Furnishings", "$15.254,37", "$-3.906,22", "-25,61%", "40,39%", "67,32%"],
            ["Central", "Tables", "$39.154,97", "$-3.559,65", "-9,09%", "26,25%", "68,06%"],
            ["Central", "Bookcases", "$24.157,18", "$-1.997,90", "-8,27%", "23,28%", "72,00%"],
            ["West", "Bookcases", "$36.004,12", "$-1.646,51", "-4,57%", "22,87%", "40,00%"],
            ["East", "Bookcases", "$43.819,33", "$-1.167,63", "-2,66%", "22,00%", "47,14%"],
            ["South", "Bookcases", "$10.899,36", "$1.339,49", "12,29%", "10,00%", "28,57%"],
            ["West", "Tables", "$84.754,56", "$1.482,61", "1,75%", "20,00%", "42,24%"],
            ["South", "Furnishings", "$17.306,68", "$3.442,68", "19,89%", "10,67%", "4,24%"],
        ],
    )

    add_spacer(doc)
    add_figure(doc, "region_sub_category_heatmap.png", "Hình 5: Heatmap lợi nhuận theo region × sub-category")

    add_spacer(doc)
    add_implication(
        doc,
        "East × Tables là tổ hợp lỗ nặng nhất (-$11.025). Central lỗ ở 3/4 sub-categories. "
        "Ưu tiên review top loss combinations và ô region × sub-category âm sâu trước."
    )

    doc.add_page_break()

    # Product scatter
    add_figure(doc, "product_sales_profit_scatter.png",
               "Hình 6: Phân bố sản phẩm theo doanh thu và lợi nhuận (đường zero-profit rõ ràng)")

    doc.add_page_break()

    # ── 4.4 Discount Risk (Group C) ──────────────────────────────
    doc.add_heading("4.4. Rủi ro discount và pricing (Nhóm C)", level=2)

    # C1: Discount band performance
    doc.add_paragraph().add_run(
        "C1. Profit và loss rate khác nhau thế nào theo discount band?"
    ).bold = True

    add_styled_table(
        doc,
        ["Discount Band", "Doanh thu", "Lợi nhuận", "Biên %", "Dòng lỗ %"],
        [
            ["0%", "$256.025,27", "$58.133,08", "22,71%", "0,00%"],
            ["1-10%", "$46.634,25", "$7.111,01", "15,25%", "5,26%"],
            ["11-20%", "$244.189,54", "$7.684,94", "3,15%", "27,44%"],
            ["21-30%", "$99.470,35", "$-10.695,32", "-10,75%", "93,24%"],
            [">30%", "$95.680,39", "$-43.782,44", "-45,76%", "100,00%"],
        ],
    )

    add_spacer(doc)
    add_figure(doc, "discount_band_margin.png", "Hình 7: Biên lợi nhuận theo discount band")

    add_spacer(doc)
    add_observation(
        doc,
        "Discount 0% biên 22,71%, tỷ lệ lỗ 0%. Discount ≥30% biên -45,76%, 100% lỗ. "
        "Nhóm discount ≥30% chiếm 26,30% doanh thu nhưng tạo lỗ -$54.477,76 "
        "— tương đương gấp 3 lần tổng lợi nhuận toàn dataset."
    )

    add_spacer(doc)

    # C2: Threshold
    doc.add_paragraph().add_run(
        "C2. Ngưỡng discount đầu tiên có aggregate profit âm:"
    ).bold = True

    doc.add_paragraph(
        "Discount 30,00% là mức đầu tiên có tổng profit âm trong bảng exact discount, "
        "với profit $-10.695,32. Nhóm discount ≥30% có profit $-54.477,76 và loss-line rate 97,23%."
    )

    add_implication(
        doc,
        "Discount ≥30% nên là exception cần phê duyệt margin, không phải công cụ bán hàng thường xuyên."
    )

    add_spacer(doc)

    # C3: High-discount exposure
    doc.add_paragraph().add_run(
        "C3. Nhóm nào phơi nhiễm discount cao nhất?"
    ).bold = True

    add_styled_table(
        doc,
        ["Sub-Category", "Region", "Doanh thu", "Lợi nhuận", "Disc %", "Loss %"],
        [
            ["Tables", "East", "$39.139,81", "$-11.025,38", "37,37%", "97,50%"],
            ["Tables", "South", "$21.381,06", "$-8.840,78", "42,04%", "100,00%"],
            ["Tables", "Central", "$22.311,33", "$-6.526,42", "37,06%", "96,08%"],
            ["Furnishings", "Central", "$6.644,70", "$-5.944,66", "60,00%", "100,00%"],
            ["Tables", "West", "$7.124,34", "$-4.305,64", "50,00%", "100,00%"],
            ["Bookcases", "East", "$7.308,46", "$-4.255,81", "50,00%", "100,00%"],
            ["Chairs", "Central", "$41.135,63", "$-4.094,34", "30,00%", "94,95%"],
            ["Bookcases", "West", "$2.459,38", "$-3.894,94", "70,00%", "100,00%"],
        ],
    )

    add_spacer(doc)
    add_figure(doc, "discount_threshold_worklist.png",
               "Hình 8: Worklist rủi ro discount — sub-category × region với discount ≥30%")

    add_spacer(doc)

    # C4: Loss without high discount
    doc.add_paragraph().add_run(
        "C4. Có pattern lỗ nào xảy ra dù không discount cao không?"
    ).bold = True

    add_styled_table(
        doc,
        ["Sub-Cat", "Region", "State", "Doanh thu", "Lợi nhuận", "Disc %", "Loss %"],
        [
            ["Tables", "West", "California", "$16.759,50", "$-1.651,02", "20%", "100%"],
            ["Chairs", "West", "California", "$12.346,45", "$-1.462,45", "20%", "100%"],
            ["Chairs", "South", "Florida", "$3.641,06", "$-478,34", "20%", "100%"],
            ["Bookcases", "East", "New York", "$11.247,98", "$-466,22", "20%", "100%"],
            ["Bookcases", "West", "California", "$5.881,15", "$-328,70", "15%", "100%"],
        ],
    )

    add_spacer(doc)
    add_hypothesis(
        doc,
        "Những dòng lỗ ở discount ≤20% gợi ý vấn đề khác ngoài discount — "
        "có thể do chi phí sản phẩm cao hoặc chi phí vận chuyển. "
        "Tuy nhiên, cần thêm dữ liệu COGS/freight để kết luận nguyên nhân."
    )

    doc.add_page_break()

    # ── 4.5 Geographic Performance (Group D) ─────────────────────
    doc.add_heading("4.5. Hiệu quả địa lý (Nhóm D)", level=2)

    # D1-D2: Region performance
    doc.add_paragraph().add_run(
        "D1-D2. Hiệu quả theo vùng (region):"
    ).bold = True

    add_styled_table(
        doc,
        ["Vùng", "Doanh thu", "Lợi nhuận", "Biên %", "Disc %", "Loss %"],
        [
            ["West", "$252.612,74", "$11.504,95", "4,55%", "13,14%", "21,92%"],
            ["East", "$208.291,20", "$3.046,17", "1,46%", "15,41%", "30,62%"],
            ["Central", "$163.797,16", "$-2.871,05", "-1,75%", "29,74%", "65,90%"],
            ["South", "$117.298,68", "$6.771,21", "5,77%", "12,15%", "17,47%"],
        ],
    )

    add_spacer(doc)
    add_observation(
        doc,
        "West biên 4,55%; South biên 5,77% — hai vùng lợi nhuận mạnh. "
        "East biên 1,46% — cần review. Central lỗ -$2.871 với discount trung bình cao 29,74% "
        "và tỷ lệ dòng lỗ 65,90% — vùng rủi ro chính."
    )

    add_spacer(doc)

    # State priority worklist
    doc.add_paragraph().add_run(
        "D2. Bảng ưu tiên tiểu bang — Top 10 state cần worklist quản trị:"
    ).bold = True

    add_styled_table(
        doc,
        ["State", "Region", "Doanh thu", "Lợi nhuận", "Biên %", "Disc %", "Loss %", "Priority"],
        [
            ["Texas", "Central", "$60.593,29", "$-10.436,14", "-17,22%", "42,30%", "97,03%", "5"],
            ["Illinois", "Central", "$28.274,52", "$-9.076,29", "-32,10%", "46,83%", "98,37%", "5"],
            ["Pennsylvania", "East", "$39.354,93", "$-7.196,72", "-18,29%", "27,68%", "52,00%", "5"],
            ["Ohio", "East", "$24.199,15", "$-4.206,32", "-17,38%", "28,49%", "56,99%", "5"],
            ["North Carolina", "South", "$15.155,48", "$-3.486,46", "-23,00%", "23,81%", "33,33%", "5"],
            ["Arizona", "West", "$13.525,29", "$-2.744,92", "-20,29%", "28,57%", "42,86%", "5"],
            ["Colorado", "West", "$13.243,04", "$-2.683,13", "-20,26%", "31,18%", "35,29%", "4"],
            ["Florida", "South", "$22.987,04", "$-2.254,98", "-9,81%", "23,24%", "32,94%", "4"],
            ["Tennessee", "South", "$13.506,73", "$-2.208,63", "-16,35%", "23,56%", "35,56%", "4"],
            ["Oregon", "West", "$6.338,13", "$-1.487,58", "-23,47%", "34,29%", "61,90%", "4"],
        ],
    )

    add_spacer(doc)
    add_figure(doc, "state_sales_profit.png",
               "Hình 9: So sánh doanh thu và lợi nhuận theo tiểu bang")

    add_spacer(doc)

    # D3-D4: State x sub-category drill-down
    doc.add_paragraph().add_run(
        "D3-D4. Drill-down: State lỗ nặng kết hợp sub-category và discount:"
    ).bold = True

    add_styled_table(
        doc,
        ["State", "Sub-Cat", "Doanh thu", "Lợi nhuận", "Biên %", "Disc %", "Loss %"],
        [
            ["New York", "Tables", "$13.779,02", "$-4.535,64", "-32,92%", "40%", "100%"],
            ["Illinois", "Tables", "$6.550,67", "$-4.309,74", "-65,79%", "50%", "100%"],
            ["N. Carolina", "Tables", "$9.681,73", "$-3.684,25", "-38,05%", "40%", "100%"],
            ["Texas", "Furnishings", "$3.766,72", "$-3.312,68", "-87,95%", "60%", "100%"],
            ["Pennsylvania", "Bookcases", "$5.230,76", "$-2.896,76", "-55,38%", "50%", "100%"],
            ["Ohio", "Tables", "$7.887,11", "$-2.715,33", "-34,43%", "40%", "100%"],
            ["Tennessee", "Tables", "$6.214,36", "$-2.663,41", "-42,86%", "40%", "100%"],
            ["Illinois", "Furnishings", "$2.877,98", "$-2.631,98", "-91,45%", "60%", "100%"],
            ["Pennsylvania", "Tables", "$8.052,19", "$-2.588,75", "-32,15%", "40%", "100%"],
            ["Texas", "Chairs", "$26.572,45", "$-2.515,65", "-9,47%", "30%", "93,44%"],
        ],
    )

    add_spacer(doc)
    add_implication(
        doc,
        "Dùng heatmap region × sub-category và state × sub-category để phân nhóm hành động: "
        "mở rộng ô xanh (lợi nhuận), bảo vệ ô profit cao, review/giảm ô âm sâu. "
        "Tất cả các ô lỗ nặng đều có discount ≥30%, xác nhận discount là yếu tố rủi ro chính."
    )

    doc.add_page_break()

    # ── 4.6 Customer & Segment (Group E) ─────────────────────────
    doc.add_heading("4.6. Khách hàng và segment (Nhóm E)", level=2)

    # E1: Segment performance
    doc.add_paragraph().add_run(
        "E1. So sánh hiệu quả giữa Consumer, Corporate, Home Office:"
    ).bold = True

    add_styled_table(
        doc,
        ["Segment", "Doanh thu", "Lợi nhuận", "Biên %", "Disc %", "Loss %"],
        [
            ["Consumer", "$391.049,31", "$6.991,08", "1,79%", "17,67%", "34,77%"],
            ["Corporate", "$229.019,79", "$7.584,82", "3,31%", "17,41%", "32,66%"],
            ["Home Office", "$121.930,70", "$3.875,38", "3,18%", "16,50%", "32,04%"],
        ],
    )

    add_spacer(doc)
    add_figure(doc, "segment_performance.png", "Hình 10: So sánh hiệu quả theo segment")

    add_spacer(doc)
    add_observation(
        doc,
        "Consumer chiếm doanh thu lớn nhất ($391.049) nhưng biên thấp nhất (1,79%) "
        "và tỷ lệ dòng lỗ cao nhất (34,77%). Corporate và Home Office có biên tốt hơn (3,18-3,31%). "
        "Segment có margin/loss-rate tốt hơn nên là nền cho cross-sell có mục tiêu."
    )

    add_spacer(doc)

    # E2: Repeat customers
    doc.add_paragraph().add_run(
        "E2. Repeat customers đóng góp bao nhiêu hoạt động?"
    ).bold = True

    add_styled_table(
        doc,
        ["Nhóm", "Khách hàng", "Đơn hàng", "Doanh thu", "Lợi nhuận", "Biên %", "% đơn hàng"],
        [
            ["Mua 1 lần", "191", "191", "$90.892,27", "$-2.108,04", "-2,32%", "10,83%"],
            ["Repeat (≥2 đơn)", "516", "1.573", "$651.107,53", "$20.559,31", "3,16%", "89,17%"],
        ],
    )

    add_spacer(doc)
    add_observation(
        doc,
        "Repeat customers chiếm 89,17% tổng số đơn trong cửa sổ dữ liệu và có biên dương (3,16%). "
        "Khách mua 1 lần có biên âm (-2,32%). Có cơ sở để ưu tiên retention/cross-sell "
        "thay vì blanket discount."
    )

    add_spacer(doc)

    # E3: High-sales low-profit customers
    doc.add_paragraph().add_run(
        "E3. Top 10 khách hàng doanh thu cao nhưng lợi nhuận âm (nhãn ẩn danh):"
    ).bold = True

    add_styled_table(
        doc,
        ["Nhãn KH", "Segment", "Đơn hàng", "Doanh thu", "Lợi nhuận", "Biên %"],
        [
            ["KH-014", "Consumer", "1", "$4.297,64", "$-1.862,31", "-43,33%"],
            ["KH-026", "Consumer", "3", "$3.631,59", "$-1.768,45", "-48,70%"],
            ["KH-030", "Corporate", "2", "$3.336,32", "$-1.125,80", "-33,74%"],
            ["KH-043", "Home Office", "4", "$2.888,11", "$-1.094,90", "-37,91%"],
            ["KH-075", "Corporate", "4", "$2.313,50", "$-1.034,12", "-44,70%"],
            ["KH-180", "Home Office", "2", "$1.418,54", "$-793,65", "-55,95%"],
            ["KH-110", "Consumer", "4", "$1.935,31", "$-730,95", "-37,77%"],
            ["KH-141", "Corporate", "3", "$1.735,56", "$-689,65", "-39,74%"],
            ["KH-153", "Consumer", "3", "$1.618,27", "$-661,80", "-40,90%"],
            ["KH-007", "Consumer", "4", "$4.899,12", "$-621,94", "-12,70%"],
        ],
    )

    add_spacer(doc)

    # E4: Cross-sell candidates
    doc.add_paragraph().add_run(
        "E4. Top 10 khách hàng lợi nhuận cao — ứng viên cross-sell:"
    ).bold = True

    add_styled_table(
        doc,
        ["Nhãn KH", "Segment", "Đơn hàng", "Doanh thu", "Lợi nhuận", "Biên %"],
        [
            ["KH-006", "Corporate", "7", "$5.387,39", "$1.146,49", "21,28%"],
            ["KH-002", "Consumer", "5", "$6.920,14", "$968,08", "13,99%"],
            ["KH-010", "Consumer", "5", "$4.513,11", "$805,98", "17,86%"],
            ["KH-009", "Corporate", "5", "$4.768,50", "$770,15", "16,15%"],
            ["KH-016", "Home Office", "5", "$4.132,06", "$750,26", "18,16%"],
            ["KH-028", "Consumer", "4", "$3.557,14", "$738,84", "20,77%"],
            ["KH-017", "Corporate", "6", "$4.078,82", "$720,58", "17,67%"],
            ["KH-001", "Consumer", "9", "$8.332,09", "$688,40", "8,26%"],
            ["KH-046", "Consumer", "3", "$2.852,39", "$688,14", "24,13%"],
            ["KH-065", "Corporate", "4", "$2.530,59", "$677,56", "26,77%"],
        ],
    )

    add_spacer(doc)
    add_implication(
        doc,
        "Khách hàng lợi nhuận cao có biên 8-27% và mua 3-9 đơn — phù hợp cho cross-sell "
        "sản phẩm Furnishings/Chairs thay vì giảm giá sâu Tables/Bookcases."
    )

    doc.add_page_break()

    # ── 4.7 Fulfillment Context (Group F) ────────────────────────
    doc.add_heading("4.7. Fulfillment context (Nhóm F)", level=2)

    doc.add_paragraph(
        "Các bảng dưới đây là mô tả quan sát, không phải kết luận chi phí vận chuyển "
        "vì dữ liệu không có freight-cost field."
    )

    # F1: Ship mode
    doc.add_paragraph().add_run(
        "F1. Hiệu quả theo phương thức vận chuyển:"
    ).bold = True

    add_styled_table(
        doc,
        ["Ship Mode", "Doanh thu", "Lợi nhuận", "Biên %", "Ngày TB", "Loss %"],
        [
            ["Standard Class", "$435.831,47", "$10.360,72", "2,38%", "4", "33,65%"],
            ["Second Class", "$156.289,02", "$4.226,26", "2,70%", "3", "32,79%"],
            ["First Class", "$110.730,52", "$3.066,95", "2,77%", "2", "33,94%"],
            ["Same Day", "$39.148,78", "$797,35", "2,04%", "0", "36,13%"],
        ],
    )

    add_spacer(doc)

    # F2: Shipping days
    doc.add_paragraph().add_run(
        "F2. Hiệu quả theo số ngày giao hàng:"
    ).bold = True

    add_styled_table(
        doc,
        ["Ngày", "Doanh thu", "Lợi nhuận", "Biên %", "Loss %"],
        [
            ["0", "$37.004,68", "$539,38", "1,46%", "36,21%"],
            ["1", "$25.113,83", "$784,14", "3,12%", "32,47%"],
            ["2", "$124.888,12", "$7.483,90", "5,99%", "32,90%"],
            ["3", "$58.159,14", "$-299,05", "-0,51%", "32,80%"],
            ["4", "$212.148,65", "$4.890,38", "2,31%", "36,61%"],
            ["5", "$161.316,71", "$137,95", "0,09%", "34,12%"],
            ["6", "$80.647,73", "$4.108,47", "5,09%", "29,66%"],
            ["7", "$42.720,92", "$806,11", "1,89%", "27,13%"],
        ],
    )

    add_spacer(doc)
    add_implication(
        doc,
        "Nếu mode hoặc số ngày giao hàng có loss-rate cao, nên xem như tín hiệu drill-down "
        "chứ không quy kết nguyên nhân vì thiếu freight-cost. Biên tương đối đồng đều (2-3%) "
        "trừ giao hàng 2 ngày (biên 5,99% — có thể do mix sản phẩm tốt hơn)."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    #  V. KHUYẾN NGHỊ HÀNH ĐỘNG
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("V. Khuyến nghị hành động", level=1)

    doc.add_paragraph(
        "Mỗi khuyến nghị dưới đây gắn với evidence trong dữ liệu và được trình bày "
        "như hành động quản trị, không phải kết luận nhân quả tuyệt đối. "
        "Mỗi khuyến nghị bao gồm: Scope, Evidence, Action, KPI, và Limitation."
    )

    add_spacer(doc)

    recommendations = [
        {
            "priority": "1",
            "scope": "Discount ≥30% (toàn dataset)",
            "evidence": "Profit $-54.477,76; loss-line rate 97,23%; chiếm 26,30% doanh thu",
            "action": "Yêu cầu review margin trước khi áp dụng giảm giá ≥30%. Chuyển từ công cụ bán hàng thường xuyên sang exception cần phê duyệt.",
            "kpi": "Giảm loss-line rate và cải thiện total profit",
            "limitation": "Không có COGS/freight nên không chứng minh nhân quả đầy đủ",
        },
        {
            "priority": "2",
            "scope": "Tables và Bookcases (sub-category)",
            "evidence": "Tables $-17.725,48 (biên -8,56%); Bookcases $-3.472,56 (biên -3,02%)",
            "action": "Review pricing, promotion và state exposure trước khi scale doanh thu. Ưu tiên sửa chữa top 10 sản phẩm lỗ lớn nhất.",
            "kpi": "Profit margin theo sub-category và top loss products",
            "limitation": "Cần thêm cost detail để tách giá vốn và phí fulfillment",
        },
        {
            "priority": "3",
            "scope": "Central và các bang ưu tiên: Texas, Illinois, Pennsylvania, Ohio",
            "evidence": "Central profit $-2.871,05; mean discount 29,74%. Texas lỗ $-10.436; Illinois lỗ $-9.076",
            "action": "Dừng blanket promotion tại điểm lỗ. Drill down theo product × discount. Áp dụng margin floor cho các bang priority 5.",
            "kpi": "State profit, discount exposure và loss-line rate",
            "limitation": "State priority score là worklist minh bạch, không phải statistical model",
        },
        {
            "priority": "4",
            "scope": "Chairs và Furnishings (sub-category lợi nhuận)",
            "evidence": "Chairs profit $26.590,17 (biên 8,10%); Furnishings margin 14,24%",
            "action": "Bảo vệ availability và test cross-sell có mục tiêu thay vì discount rộng. Scale Furnishings selectively.",
            "kpi": "Sales growth, profit growth và weighted margin",
            "limitation": "Không có tồn kho nên không kết luận được stockout",
        },
        {
            "priority": "5",
            "scope": "Repeat customers và segments",
            "evidence": "Repeat-order share 89,17%; repeat customers biên 3,16% vs mua 1 lần biên -2,32%",
            "action": "Ưu tiên offer có mục tiêu cho khách lặp lại/có lợi nhuận thay vì giảm giá đại trà. Cross-sell Furnishings cho khách Corporate biên cao.",
            "kpi": "Profit per order, AOV và customer margin",
            "limitation": "Chỉ là within-window behavior trong giai đoạn 2014-2017",
        },
    ]

    add_styled_table(
        doc,
        ["#", "Scope", "Evidence", "Action", "KPI", "Limitation"],
        [
            [r["priority"], r["scope"], r["evidence"], r["action"], r["kpi"], r["limitation"]]
            for r in recommendations
        ],
    )

    add_spacer(doc)

    # Detailed recommendation blocks
    for r in recommendations:
        p = doc.add_paragraph()
        run = p.add_run(f"Khuyến nghị #{r['priority']}: {r['scope']}")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

        details = [
            ("Scope: ", r["scope"]),
            ("Evidence: ", r["evidence"]),
            ("Action: ", r["action"]),
            ("KPI theo dõi: ", r["kpi"]),
            ("Hạn chế: ", r["limitation"]),
        ]
        for label, val in details:
            bp = doc.add_paragraph(style="List Bullet")
            rl = bp.add_run(label)
            rl.bold = True
            rl.font.size = Pt(10)
            rv = bp.add_run(val)
            rv.font.size = Pt(10)

        add_spacer(doc)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    #  VI. TỔNG QUAN VÀ KẾT LUẬN
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("VI. Tổng quan và kết luận", level=1)

    doc.add_paragraph(
        "Qua quá trình phân tích khám phá bộ dữ liệu bán hàng Furniture, "
        "nhóm rút ra các kết quả chính sau:"
    )

    conclusions = [
        (
            "1. Doanh thu và lợi nhuận phân kỳ: ",
            "Doanh thu $741.999,80 nhưng lợi nhuận chỉ $18.451,27 (biên 2,49%), "
            "với 33,66% dòng lỗ. Chất lượng doanh số là vấn đề cấu trúc."
        ),
        (
            "2. Sự không đồng bộ giữa các năm: ",
            "Doanh thu tăng 37% từ 2014-2017 nhưng lợi nhuận không tăng theo tỷ lệ. "
            "Năm 2017 doanh thu cao nhất nhưng lợi nhuận thấp nhất (1,40% biên)."
        ),
        (
            "3. Sub-category không cân bằng: ",
            "Chairs lợi nhuận +$26.590; Furnishings biên 14,24%; Tables lỗ -$17.725; "
            "Bookcases lỗ -$3.473. Chiến lược cần bảo vệ nhóm lợi nhuận và sửa nhóm lỗ."
        ),
        (
            "4. Discount là yếu tố rủi ro chính: ",
            "Discount 0% biên 22,71%; discount ≥30% biên -45,76% và 100% lỗ. "
            "Nhóm discount ≥30% chiếm 26,30% doanh thu nhưng lỗ $-54.477,76."
        ),
        (
            "5. Rủi ro địa lý tập trung: ",
            "Central lỗ $-2.871 với discount trung bình 29,74%. "
            "Texas (-$10.436), Illinois (-$9.076), Pennsylvania (-$7.197) là điểm nóng. "
            "West và South là vùng lợi nhuận mạnh."
        ),
        (
            "6. Repeat customers là tài sản: ",
            "89,17% đơn hàng đến từ khách lặp lại, biên dương 3,16%. "
            "Khách mua 1 lần có biên âm (-2,32%). Retention quan trọng hơn acquisition."
        ),
    ]

    for label, text in conclusions:
        p = doc.add_paragraph()
        r = p.add_run(label)
        r.bold = True
        r.font.size = Pt(11)
        r2 = p.add_run(text)
        r2.font.size = Pt(11)

    add_spacer(doc)

    # Final narrative
    p = doc.add_paragraph()
    p.style = "Normal"
    r = p.add_run("Narrative tổng kết: ")
    r.bold = True
    r.font.size = Pt(11)
    r2 = p.add_run(
        "Doanh thu Furniture tăng trưởng trong giai đoạn quan sát, nhưng lợi nhuận "
        "không scale theo doanh thu. Vấn đề quản trị trung tâm không phải là tạo cầu "
        "(demand generation) mà là chất lượng lợi nhuận (profit quality). "
        "Tables và Bookcases, các giao dịch discount cao, và một số bang doanh thu lớn "
        "là nguồn lỗ chính. Chairs và Furnishings cung cấp con đường đáng tin cậy nhất "
        "cho tăng trưởng có lợi nhuận, trong khi discount ≥30% nên được coi là exception "
        "kiểm soát biên chứ không phải đòn bẩy thường xuyên."
    )
    r2.font.size = Pt(11)

    add_spacer(doc)
    doc.add_paragraph(
        "Những phát hiện này cung cấp cơ sở rõ ràng để hướng dẫn các quyết định chiến lược "
        "về sản phẩm, giá cả, và khuyến mại trong danh mục Furniture, nhằm mục tiêu "
        "tăng doanh thu bền vững mà không hy sinh lợi nhuận."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    #  VII. PHỤ LỤC
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("VII. Phụ lục", level=1)

    doc.add_heading("7.1. Danh sách biến phái sinh (derived fields)", level=2)

    add_styled_table(
        doc,
        ["Trường mới", "Định nghĩa", "Loại"],
        [
            ["order_year", "Năm của Order Date", "integer"],
            ["order_quarter", "Quý: YYYY-QN", "string"],
            ["order_month", "Tháng bắt đầu", "datetime"],
            ["order_month_label", "Label tháng: YYYY-MM", "string"],
            ["order_weekday", "Ngày trong tuần", "ordered category"],
            ["shipping_days", "Số ngày giao hàng = Ship Date − Order Date", "integer"],
            ["profit_margin", "profit / sales (chỉ khi sales > 0)", "float"],
            ["is_loss_line", "True nếu profit < 0", "boolean"],
            ["discount_pct", "discount × 100", "float"],
            ["discount_band", "0% / 1-10% / 11-20% / 21-30% / >30%", "ordered category"],
            ["sales_per_unit", "sales / quantity", "float"],
            ["profit_per_unit", "profit / quantity", "float"],
            ["is_repeat_customer", "True nếu khách có ≥ 2 đơn", "boolean"],
            ["customer_order_count", "nunique(Order ID) theo Customer ID", "integer"],
            ["customer_sales", "Tổng sales theo khách hàng", "float"],
            ["customer_profit", "Tổng profit theo khách hàng", "float"],
            ["customer_recency_days", "Ngày kể từ đơn hàng cuối đến cutoff", "integer"],
        ],
    )

    add_spacer(doc)

    doc.add_heading("7.2. Discount bands — định nghĩa chính thức", level=2)

    add_styled_table(
        doc,
        ["Label", "Cận dưới (inclusive)", "Cận trên (inclusive)"],
        [
            ["0%", "0,00", "0,00"],
            ["1-10%", ">0,00", "0,10"],
            ["11-20%", ">0,10", "0,20"],
            ["21-30%", ">0,20", "0,30"],
            [">30%", ">0,30", "1,00"],
        ],
    )

    add_spacer(doc)

    doc.add_heading("7.3. State priority score — công thức", level=2)

    doc.add_paragraph("Priority score là triage mechanism minh bạch, không phải statistical model:")

    score_items = [
        "+1 điểm nếu tổng profit < 0",
        "+1 điểm nếu doanh thu > median doanh thu các state",
        "+1 điểm nếu mean discount > 20%",
        "+1 điểm nếu loss-line rate > 50%",
        "+1 điểm nếu absolute loss > $2.000",
    ]
    for s in score_items:
        doc.add_paragraph(s, style="List Bullet")

    add_spacer(doc)

    doc.add_heading("7.4. Product classification matrix", level=2)

    add_styled_table(
        doc,
        ["Phân loại", "Quy tắc", "Hành động đề xuất"],
        [
            ["Grow", "Doanh thu cao, profit dương, biên chấp nhận được", "Bảo vệ availability, test tăng trưởng có mục tiêu"],
            ["Repair", "Doanh thu cao, profit thấp/âm", "Review giá, discount, chi phí sản phẩm, state exposure"],
            ["Scale Selectively", "Doanh thu thấp/vừa, biên cao", "Cải thiện discoverability, cross-sell"],
            ["Rationalize", "Doanh thu thấp, profit thấp/âm", "Review giảm promotion, thay đổi giá, hoặc ngừng"],
        ],
    )

    add_spacer(doc)

    doc.add_heading("7.5. Quarterly performance đầy đủ", level=2)

    add_styled_table(
        doc,
        ["Quý", "Doanh thu", "Lợi nhuận", "Biên %", "Tăng DT", "Tăng LN"],
        [
            ["2014-Q1", "$22.656,14", "$-202,50", "-0,89%", "—", "—"],
            ["2014-Q2", "$28.063,75", "$800,82", "2,85%", "23,87%", "495,47%"],
            ["2014-Q3", "$41.957,88", "$2.896,32", "6,90%", "49,51%", "261,67%"],
            ["2014-Q4", "$64.515,09", "$1.963,09", "3,04%", "53,76%", "-32,22%"],
            ["2015-Q1", "$27.374,10", "$-1.164,25", "-4,25%", "-57,57%", "-159,31%"],
            ["2015-Q2", "$27.564,83", "$826,58", "3,00%", "0,70%", "171,00%"],
            ["2015-Q3", "$49.586,04", "$537,55", "1,08%", "79,89%", "-34,97%"],
            ["2015-Q4", "$65.993,28", "$2.815,32", "4,27%", "33,09%", "423,73%"],
        ],
    )

    add_spacer(doc)

    doc.add_heading("7.6. Danh sách artifacts đã tạo", level=2)

    artifacts = [
        "data/processed/furniture_sales_clean.parquet — Bảng fact đã làm sạch",
        "data/processed/customer_summary.parquet — Tóm tắt theo khách hàng",
        "data/processed/product_summary.parquet — Tóm tắt theo sản phẩm",
        "data/processed/order_summary.parquet — Tóm tắt theo đơn hàng",
        "data/processed/metric_baseline.json — Các giá trị baseline và metadata",
        "reports/data_quality_report.md — Báo cáo chất lượng dữ liệu",
        "reports/executive_summary.md — Tóm tắt điều hành",
        "reports/business_questions.md — Trả lời câu hỏi kinh doanh A-F",
        "reports/recommendations.md — Khuyến nghị hành động",
        "reports/tables/*.csv — 34 bảng drill-down chi tiết",
        "reports/figures/*.png — 10 biểu đồ chuyên nghiệp",
    ]
    for a in artifacts:
        doc.add_paragraph(a, style="List Bullet")

    add_spacer(doc)

    doc.add_heading("7.7. Hạn chế và phạm vi phân tích", level=2)

    limitations = [
        "Phân tích chỉ bao phủ danh mục Furniture — không mở rộng kết luận cho toàn công ty.",
        "Không có COGS (giá vốn hàng bán) — không thể tái tạo gross margin chính xác.",
        "Không có freight cost, marketing spend, returns — không xác nhận cơ chế nhân quả đầy đủ.",
        "Discount và profit có tương quan mạnh, nhưng tương quan ≠ nhân quả.",
        "Customer metrics là within-window (2014-2017), không phải lifetime behavior.",
        "Forecast (nếu có) chỉ mang tính khám phá, cần temporal validation riêng.",
    ]
    for l in limitations:
        doc.add_paragraph(l, style="List Bullet")

    # ── Save ───────────────────────────────────────────────────────
    doc.save(str(OUTPUT))
    print(f"✅ Report saved to: {OUTPUT}")
    print(f"   File size: {OUTPUT.stat().st_size:,} bytes")


if __name__ == "__main__":
    generate_report()
