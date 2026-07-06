import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Inches

doc = Document('MauBaoCao.docx')

print('=== STYLES USED ===')
styles = set()
for p in doc.paragraphs:
    styles.add(p.style.name)
print(sorted(styles))

print('\n=== TABLE STYLES ===')
for i, t in enumerate(doc.tables):
    print(f'Table {i}: style={t.style.name}')

print('\n=== FONT INFO FROM RUNS ===')
for p in doc.paragraphs[:10]:
    for r in p.runs:
        print(f'Style={p.style.name}, Text={r.text[:50]}, Font={r.font.name}, Size={r.font.size}, Bold={r.font.bold}')

print('\n=== SECTION LAYOUT ===')
for i, s in enumerate(doc.sections):
    print(f'Section {i}: width={s.page_width}, height={s.page_height}')
    print(f'  margins: top={s.top_margin}, bottom={s.bottom_margin}, left={s.left_margin}, right={s.right_margin}')

print('\n=== TABLE 0 CONTENT ===')
t = doc.tables[0]
for row in t.rows:
    print(' | '.join([c.text[:50] for c in row.cells]))
